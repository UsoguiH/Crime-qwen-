"""DOCX export via docxtpl from a programmatically-built RTL template.

The template is regenerated on every render (cheap, never stale) with
python-docx + oxml surgery for w:bidi / w:rtl / w:rFonts@cs / w:szCs, so every
`{{ var }}` substitution inherits correct Arabic styling and Word shapes the
text natively. Control flow uses PLAIN Jinja tags in dedicated paragraphs —
the fragment between them forms closed XML for any iteration count, which
avoids docxtpl's `{%p/%tr` structural preprocessing entirely (its only cost is
an empty paragraph where each tag lived).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ARABIC_FONT = "IBM Plex Sans Arabic"
FALLBACK_FONT = "Arial"


def _set_rtl_paragraph(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _style_run(run, size_pt: int = 11, bold: bool = False) -> None:
    run.font.name = FALLBACK_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.rtl = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), ARABIC_FONT)
    rFonts.set(qn("w:ascii"), FALLBACK_FONT)
    rFonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(size_pt * 2))
    if bold and rPr.find(qn("w:bCs")) is None:
        rPr.append(OxmlElement("w:bCs"))


def _p(doc, text: str, size: int = 11, bold: bool = False):
    paragraph = doc.add_paragraph()
    _set_rtl_paragraph(paragraph)
    run = paragraph.add_run(text)
    _style_run(run, size, bold)
    return paragraph


def build_template(dst: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FALLBACK_FONT
    normal.font.size = Pt(11)
    normal.font.rtl = True

    _p(doc, "سري", 14, bold=True)
    _p(doc, "تقرير تحليلي لمسرح الجريمة — نظام أثر", 18, bold=True)
    _p(doc, "رقم القضية: {{ case.case_number_ar }}", 12)
    _p(doc, "العنوان: {{ case.title_ar }}", 12)
    _p(doc, "الموقع: {{ case.location_ar }}", 12)
    _p(doc, "المحقق: {{ case.investigator_name_ar }}", 12)
    _p(doc, "تاريخ الواقعة: {{ case.incident_dual_date or 'غير محدد' }}", 12)
    _p(doc, "تاريخ إصدار التقرير: {{ generated.dual_date }}", 12)

    _p(doc, "أولاً: الملخص التنفيذي", 14, bold=True)
    _p(doc, "{{ narratives.exec_summary or 'لا يتوفر.' }}")

    _p(doc, "ثانياً: السرد الزمني للأحداث", 14, bold=True)
    _p(doc, "{{ narratives.timeline or 'لا يتوفر.' }}")
    _p(doc, "{% for ev in events %}")
    _p(doc, "— {{ ev.text }}")
    _p(doc, "{% endfor %}")

    _p(doc, "ثالثاً: جدول الأدلة", 14, bold=True)
    _p(doc, "{% for e in entities %}")
    _p(doc, "{{ e.label_ar }} — {{ e.name_ar }} — {{ e.category_ar }} — "
            "ثقة النموذج: {{ e.confidence_pct }} — {{ e.review_status_ar }}")
    _p(doc, "{% endfor %}")

    _p(doc, "رابعاً: التحليل التفصيلي للأدلة", 14, bold=True)
    _p(doc, "{% for e in entities %}")
    _p(doc, "{{ e.label_ar }} — {{ e.name_ar }} ({{ e.category_ar }})", 12, bold=True)
    _p(doc, "الوصف: {{ e.description_ar }}")
    _p(doc, "الدلالة الجنائية المحتملة: {{ e.forensic_significance_ar }}")
    _p(doc, "توصية التعامل: {{ e.handling_recommendation_ar }}")
    _p(doc, "المصادر: {{ e.sources|join('، ') }} — درجة ثقة النموذج: {{ e.confidence_pct }}"
            "{% if e.needs_review %} — يتطلب مراجعة بشرية{% endif %}")
    _p(doc, "{% endfor %}")

    _p(doc, "خامساً: العلاقات المكانية", 14, bold=True)
    _p(doc, "{{ narratives.spatial or 'لا يتوفر.' }}")
    _p(doc, "سادساً: ما يتطلب تحقيقاً أعمق ومراجعة بشرية", 14, bold=True)
    _p(doc, "{{ narratives.review_needed or 'لا يتوفر.' }}")
    _p(doc, "سابعاً: التوصيات", 14, bold=True)
    _p(doc, "{{ narratives.recommendations or 'لا يتوفر.' }}")

    _p(doc, "ثامناً: سلسلة الحيازة", 14, bold=True)
    _p(doc, "{% for m in media %}")
    _p(doc, "• {{ m.label }} ({{ m.kind_ar }}) — SHA-256: {{ m.sha256 }}")
    _p(doc, "{% endfor %}")
    _p(doc, "رأس سلسلة سجل التدقيق: {{ audit_head }}")

    _p(doc, "{{ disclaimer }}", 11, bold=True)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def render_docx(template_path: Path, context: dict, dst: Path) -> None:
    from docxtpl import DocxTemplate

    build_template(template_path)  # always regenerate: cheap, never stale
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    dst.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(dst))
